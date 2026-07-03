import base64
import io

from docx import Document
from docx.shared import Inches
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

from ..schemas.configurator import DigitalTwinResponse
from .spec_text_gen import (
    build_spec_context,
    get_communication_subject,
    get_enclosure_clauses,
    get_safe_control_philosophy_clauses,
    get_starter_templates,
)

# Source - https://stackoverflow.com/a/68530806
# Posted by JGC
# Retrieved 2026-05-29, License - CC BY-SA 4.0


# Stack overflow goated
def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    pPr.insert_element_before(
        pBdr,
        "w:shd",
        "w:tabs",
        "w:suppressAutoHyphens",
        "w:kinsoku",
        "w:wordWrap",
        "w:overflowPunct",
        "w:topLinePunct",
        "w:autoSpaceDE",
        "w:autoSpaceDN",
        "w:bidi",
        "w:adjustRightInd",
        "w:snapToGrid",
        "w:spacing",
        "w:ind",
        "w:contextualSpacing",
        "w:mirrorIndents",
        "w:suppressOverlap",
        "w:jc",
        "w:textDirection",
        "w:textAlignment",
        "w:textboxTightWrap",
        "w:outlineLvl",
        "w:divId",
        "w:cnfStyle",
        "w:rPr",
        "w:sectPr",
        "w:pPrChange",
    )
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)


def generate_word_from_twin(twin: DigitalTwinResponse) -> bytes:
    doc = Document()

    ctx = build_spec_context(twin)
    has_plc = ctx["has_plc"]
    has_scada = ctx["has_scada"]
    has_network = ctx["has_network"]
    comm_mode = ctx["comm_mode"]
    comm_label = ctx["comm_label"]
    starter_kind = ctx["starter_kind"]
    device_type = ctx["device_type"]
    supports_speed_reference = ctx["supports_speed_reference"]

    enclosure = getattr(twin, "enclosure", None)
    mounting = getattr(enclosure, "mounting_type", "Floor Standing")
    dimensions = getattr(enclosure, "dimensions_mm", "N/A")
    ip_rating = getattr(enclosure, "ip_rating", "IP rating TBC")

    if starter_kind == "vfd":
        control_summary = "Pressure control"
    else:
        control_summary = "Pump staging control"

    comm_summary = comm_label if has_network else "Hardwired"

    doc.add_heading("Application Specification Appendix", 0)
    doc.add_paragraph(f"Project Configuration DNA: {twin.config_id}")
    doc.add_paragraph(f"Water Booster Set - {twin.load_count}-pump system")

    if getattr(twin, "project_name", None):
        doc.add_paragraph(f"Project Name: {twin.project_name}")
    if getattr(twin, "project_client", None):
        doc.add_paragraph(f"Customer: {twin.project_client}")
    if getattr(twin, "project_technical_manager", None):
        doc.add_paragraph(f"Technical Manager: {twin.project_technical_manager}")
    if getattr(twin, "project_location", None):
        doc.add_paragraph(f"Location: {twin.project_location}")
    if getattr(twin, "project_date", None):
        doc.add_paragraph(f"Date: {twin.project_date}")
    if getattr(twin, "project_notes", None):
        doc.add_paragraph(f"Additional Notes: {twin.project_notes}")

    insertHR(doc.add_paragraph())

    doc.add_paragraph(
        f"{device_type} Control Panel - "
        f"{twin.load_count} x {twin.motor_power_kw} kW | "
        f"{dimensions} | {ip_rating} {mounting} | "
        f"{control_summary} | {comm_summary}"
    )

    # 2. Application Data Sheet
    doc.add_heading("1. Application Data Sheet", level=1)
    table = doc.add_table(rows=7, cols=2)
    table.style = "Table Grid"

    def add_row(idx, key, value):
        table.rows[idx].cells[0].text = key
        table.rows[idx].cells[1].text = value

    if starter_kind == "vfd":
        process_objective = (
            "Maintain discharge pressure at setpoint across variable demand"
        )
        control_mode = (
            "Pressure control with pump staging/de-staging and automatic alternation "
            "where provided by the selected controller"
        )
    elif starter_kind in ["ats01", "ats130"]:
        process_objective = "Maintain discharge pressure within the required operating band across variable demand"
        control_mode = (
            "Pressure-based pump staging/de-staging with automatic alternation using "
            "soft starter start/stop control where provided by the selected controller"
        )
    elif starter_kind == "dol_adv":
        process_objective = "Maintain discharge pressure within the required operating band across variable demand"
        control_mode = (
            "Pressure-based pump staging/de-staging with automatic alternation using "
            "DOL start/stop control with advanced motor management where provided by the selected controller"
        )
    elif starter_kind == "dol":
        process_objective = "Maintain discharge pressure within the required operating band across variable demand"
        control_mode = (
            "Pressure-based pump staging/de-staging with automatic alternation using "
            "DOL start/stop control where provided by the selected controller"
        )
    else:
        process_objective = "Maintain discharge pressure according to the selected booster control philosophy"
        control_mode = "Pressure-based pump staging/de-staging with automatic alternation where provided by the selected controller"

    if has_network:
        controller_side = (
            "PLC" if has_plc else "external controller or site control system"
        )
        comm_row = (
            f"{comm_label} interface for connection between the {controller_side} "
            "and selected communication-capable devices"
        )
    else:
        comm_row = (
            "Hardwired control and monitoring signals between the controller interface "
            "and motor starters"
        )

    if has_scada:
        if has_plc:
            comm_row += "; SCADA interface included"
        else:
            comm_row += "; SCADA integration via external controller, BMS, or site control system"

    add_row(0, "Application", f"Water Booster Set - {twin.load_count} Pumps")
    add_row(1, "Process objective", process_objective)
    add_row(
        2,
        "Motor configuration",
        f"{twin.load_count} pumps x {twin.motor_power_kw} kW each",
    )
    add_row(3, "Starter type", f"{device_type} for each pump")
    add_row(4, "Control mode", control_mode)
    add_row(5, "Enclosure", f"{ip_rating} {mounting.lower()} control panel")
    add_row(6, "Communications", comm_row)

    # 2.1 Technical Description
    if twin.series_id == "DOL":
        from ..utils.ui_desc_func_prot_txt.dol import (
            DESCRIPTION,
            TECHNICAL_CHARACTERISTICS,
            FUNCTIONS,
            PROTECTIONS,
        )
    elif twin.series_id == "DOL_ADV":
        from ..utils.ui_desc_func_prot_txt.dola import (
            DESCRIPTION,
            TECHNICAL_CHARACTERISTICS,
            FUNCTIONS,
            PROTECTIONS,
        )
    elif twin.series_id == "SS":
        if float(twin.motor_power_kw) <= 15:
            from ..utils.ui_desc_func_prot_txt.ats01 import (
                DESCRIPTION,
                TECHNICAL_CHARACTERISTICS,
                FUNCTIONS,
                PROTECTIONS,
            )
        else:
            from ..utils.ui_desc_func_prot_txt.ats130 import (
                DESCRIPTION,
                TECHNICAL_CHARACTERISTICS,
                FUNCTIONS,
                PROTECTIONS,
            )
    elif twin.series_id == "VSD":
        from ..utils.ui_desc_func_prot_txt.vsd import (
            DESCRIPTION,
            TECHNICAL_CHARACTERISTICS,
            FUNCTIONS,
            PROTECTIONS,
        )
    else:
        from ..utils.ui_desc_func_prot_txt.ats01 import (
            DESCRIPTION,
            TECHNICAL_CHARACTERISTICS,
            FUNCTIONS,
            PROTECTIONS,
        )

    from ..utils.ui_desc import extract_enclosure_range

    mounting_type = twin.enclosure.mounting_type if twin.enclosure else "wall-mounted"
    catalog_ref = twin.enclosure.catalog_ref if twin.enclosure else "dummy"
    enclosure_range = extract_enclosure_range(catalog_ref)
    material = (
        "Polyester"
        if ("PLA" in catalog_ref.upper() or "PLM" in catalog_ref.upper())
        else "Steel"
    )

    tech_chars = TECHNICAL_CHARACTERISTICS
    tech_chars = tech_chars.replace("{{PUMP_COUNT}}", str(twin.load_count))
    tech_chars = tech_chars.replace("{{MOUNTING_TYPE}}", mounting_type)
    tech_chars = tech_chars.replace("{{MATERIAL}}", material.lower())
    tech_chars = tech_chars.replace("{{RANGE}}", enclosure_range)

    doc.add_heading("2. Technical Description", level=1)
    doc.add_heading("Description", level=2)
    doc.add_paragraph(DESCRIPTION)
    doc.add_heading("Technical Characteristics", level=2)
    doc.add_paragraph(tech_chars)
    doc.add_heading("Functions", level=2)
    doc.add_paragraph(FUNCTIONS)
    doc.add_heading("Protections", level=2)
    doc.add_paragraph(PROTECTIONS)

    # 3. Multi-Line Diagram & Reference Architecture
    doc.add_heading("3. Multi-Line Diagram & Reference Architecture", level=1)
    doc.add_paragraph(
        "The multi-line diagram below provides a typical control and power arrangement for a multi-pump booster set. It is intended as an early-stage template."
    )

    if hasattr(twin, "multi_line_diagram_b64") and twin.multi_line_diagram_b64:
        try:
            b64_data = twin.multi_line_diagram_b64
            if "," in b64_data:
                b64_data = b64_data.split(",")[1]
            image_data = base64.b64decode(b64_data)
            image_stream = io.BytesIO(image_data)
            doc.add_picture(image_stream, width=Inches(6.0))
            doc.add_paragraph(
                "Figure - Generated multi line diagram injected from digital twin configuration."
            )
        except Exception as e:
            doc.add_paragraph(f"[Image Generation Failed: {e}]")
    else:
        doc.add_paragraph(
            "[PLACEHOLDER: Generated multi line diagram will be inserted here when requested via UI payload.]"
        )

    doc.add_paragraph(
        f"The reference architecture below provides a typical power and control arrangement for a multi-pump booster set with {device_type} control. It is intended as an early-stage template. The detailed GA, heat dissipation, cable entry, and assembly details shall be provided by the selected panel builder."
    )

    if hasattr(twin, "reference_architecture_b64") and twin.reference_architecture_b64:
        try:
            b64_data = twin.reference_architecture_b64
            if "," in b64_data:
                b64_data = b64_data.split(",")[1]
            image_data = base64.b64decode(b64_data)
            image_stream = io.BytesIO(image_data)
            doc.add_picture(image_stream, width=Inches(6.0))
            doc.add_paragraph(
                "Figure - Generated reference architecture diagram injected from digital twin configuration."
            )
        except Exception as e:
            doc.add_paragraph(f"[Image Generation Failed: {e}]")
    else:
        doc.add_paragraph(
            "[PLACEHOLDER: Generated reference architecture diagram will be inserted here when requested via UI payload.]"
        )

    # 3.1 Bill of Materials
    import os
    import json
    row_index = 1
    bom_data = []
    json_path = os.path.join(
        os.path.dirname(__file__),
        "added_hardware",
        "00_added_hardware_and_mode_logic.json",
    )
    try:
        with open(json_path, "r") as f:
            hardware_logic = json.load(f)
        hardware_map = {item["Circuit breaker"]: item for item in hardware_logic}
    except FileNotFoundError:
        hardware_map = {}

    if twin.bom_lines:
        for line in twin.bom_lines:
            bom_data.append(
                {
                    "Index": str(row_index),
                    "Item Category": line.item_category,
                    "Item": line.item,
                    "Qty": float(line.qty),
                    "Part No.": line.part_number,
                    "Key Selection Notes / Options": (
                        line.key_selection_notes or "-"
                    ),
                }
            )
            row_index += 1

            # Apply hardware addition logic
            if line.part_number and line.part_number in hardware_map:
                hw = hardware_map[line.part_number]
                qty_per_cb = hw["Recommended qty per CB"]

                # Add auxiliary
                bom_data.append(
                    {
                        "Index": str(row_index),
                        "Item Category": "Auxiliary",
                        "Item": hw["OF auxiliary"],
                        "Qty": float(line.qty) * qty_per_cb,
                        "Part No.": hw["OF auxiliary"],
                        "Key Selection Notes / Options": hw["Notes"],
                    }
                )
                row_index += 1

    doc.add_heading("4. Bill of Materials", level=1)
    doc.add_paragraph("The table below lists the components configured for this booster set control panel:")

    bom_table = doc.add_table(rows=1 + len(bom_data), cols=6)
    bom_table.style = "Table Grid"
    bom_table.allow_autofit = False

    # Set column widths
    from docx.shared import Pt
    col_widths = [Inches(0.5), Inches(1.1), Inches(1.5), Inches(0.5), Inches(1.3), Inches(1.6)]
    for idx, width in enumerate(col_widths):
        bom_table.columns[idx].width = width
    for row in bom_table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    headers = ["Index", "Item Category", "Item", "Qty", "Part No.", "Key Selection Notes / Options"]
    for i, h in enumerate(headers):
        cell = bom_table.rows[0].cells[i]
        cell.text = h
        # Make headers bold and size 9.5
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)

    for r_idx, item in enumerate(bom_data):
        row = bom_table.rows[1 + r_idx]
        
        # Populate text
        row.cells[0].text = str(item["Index"])
        row.cells[1].text = str(item["Item Category"])
        row.cells[2].text = str(item["Item"])
        row.cells[3].text = str(item["Qty"])
        row.cells[4].text = str(item["Part No."])
        row.cells[5].text = str(item["Key Selection Notes / Options"])

        # Format font size for Item Category (1), Item (2), and Key Selection Notes / Options (5) to Pt(8.5)
        small_cols = [1, 2, 4, 5]
        for col_idx in small_cols:
            cell = row.cells[col_idx]
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
                    
        # Other columns Pt(9.5)
        other_cols = [0, 3, 4]
        for col_idx in other_cols:
            cell = row.cells[col_idx]
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)

    # 4. Panel Scope and Deliverables
    doc.add_heading("5. Panel Scope and Deliverables", level=1)
    doc.add_paragraph("The booster control panel assembly shall include, as a minimum:")

    doc.add_paragraph(
        "Main incoming isolator and/or MCCB sized for the panel maximum demand and fault level.",
        style="List Bullet",
    )

    if starter_kind == "vfd":
        feeder_label = "VSD feeders"
    elif starter_kind in ["ats01", "ats130"]:
        feeder_label = "soft starter feeders"
    elif starter_kind == "dol_adv":
        feeder_label = "DOL starter feeders with advanced motor management"
    elif starter_kind == "dol":
        feeder_label = "DOL starter feeders"
    else:
        feeder_label = "motor starter feeders"

    doc.add_paragraph(
        f"{twin.load_count} {feeder_label} sized for {twin.motor_power_kw} kW motors, including appropriate upstream protection, isolation, and motor protection as per applicable standards.",
        style="List Bullet",
    )

    if has_plc:
        doc.add_paragraph(
            "PLC with sufficient I/O and processing capacity for pressure feedback, permissives, local controls, alarms, and pump sequencing.",
            style="List Bullet",
        )
    else:
        if comm_mode == "hardwired":
            doc.add_paragraph(
                "Hardwired interface terminals for connection to the external pump controller, BMS, or site control system.",
                style="List Bullet",
            )
        else:
            doc.add_paragraph(
                "Hardwired and/or communication interface terminals for connection to the external pump controller, BMS, or site control system.",
                style="List Bullet",
            )

    doc.add_paragraph(
        "Local hardwired operator devices as required, such as selector switches, pushbuttons, reset pushbuttons, and pilot lights.",
        style="List Bullet",
    )

    if has_network:
        doc.add_paragraph(
            f"{comm_label} communication interface provisions for connection to the selected controller or site control system.",
            style="List Bullet",
        )
    else:
        doc.add_paragraph(
            "Hardwired terminal interface for pump commands, feedback signals, permissives, and alarms.",
            style="List Bullet",
        )

    doc.add_paragraph(
        "Panel auxiliaries: control power supplies where required, control protection, terminal blocks, labeling, earthing, and service accessories such as lighting or heater where required.",
        style="List Bullet",
    )

    deliverables = [
        "electrical drawings",
        "single-line diagram",
        "control schematics",
        "general arrangement drawing",
        "terminal schedule",
        "BOM",
        "FAT/SAT records",
        "operating and maintenance manuals",
    ]

    if has_plc:
        deliverables.extend(["I/O list", "PLC program backup"])

    if comm_mode == "hardwired":
        deliverables.append("hardwired interface schedule")
    else:
        deliverables.extend(
            [
                "network architecture",
                "addressing plan where applicable",
                "communication parameter list",
                "available data point list",
            ]
        )

    if has_scada:
        deliverables.append("SCADA interface description")

    doc.add_paragraph(
        "Documentation set shall include: " + ", ".join(deliverables) + ".",
        style="List Bullet",
    )

    # 6. Technical Requirements
    doc.add_heading("6. Technical Requirements", level=1)

    doc.add_heading("6.1 Enclosure and Assembly", level=2)
    doc.add_paragraph(
        f"The control panel enclosure shall be {mounting.lower()} with minimum ingress protection rating {ip_rating} and suitable for indoor technical room installation unless stated otherwise."
    )

    enc_lines = get_enclosure_clauses(twin)
    for line in enc_lines:
        doc.add_paragraph(line)

    doc.add_heading("6.2 Feeders", level=2)
    doc.add_paragraph(
        f"The panel shall include {twin.load_count} motor feeder sections, one per pump."
    )
    doc.add_paragraph(
        "The design shall include appropriate upstream short-circuit protection, overload protection, isolation, and earthing for each feeder."
    )

    templates = get_starter_templates(twin)
    feeder_clauses = [
        clause.format(motor_power_kw=twin.motor_power_kw)
        for clause in templates["feeders"]
    ]

    for clause in feeder_clauses:
        doc.add_paragraph(clause)

    if comm_mode == "hardwired":
        if starter_kind == "vfd":
            doc.add_paragraph(
                f"Each {device_type} shall support hardwired I/O for start/stop, speed reference, run feedback, fault feedback, and monitoring."
            )
        elif starter_kind == "dol_adv":
            doc.add_paragraph(
                "Each DOL starter with advanced motor management shall provide hardwired interfaces for start command, stop command or run permissive, run feedback, fault feedback, motor-management trip indication, reset where applicable, diagnostics, and monitoring."
            )
        elif starter_kind == "dol":
            doc.add_paragraph(
                "Each DOL starter shall provide hardwired interfaces for start command, stop command or run permissive, run feedback, fault feedback, overload trip indication, reset where applicable, and monitoring."
            )
        else:
            doc.add_paragraph(
                f"Each {device_type} shall support hardwired I/O for start command, stop command or run permissive, run feedback, fault feedback, reset where applicable, and monitoring."
            )
    else:
        comm_subject = get_communication_subject(ctx)

        if starter_kind == "vfd":
            doc.add_paragraph(
                f"{comm_subject} shall support {comm_label} communications for start/stop, speed reference, status, diagnostics, and monitoring."
            )
        elif starter_kind == "dol_adv":
            doc.add_paragraph(
                f"{comm_subject} shall provide {comm_label}-capable interfaces for pump start/stop commands, run feedback, ready status, fault status, motor-management diagnostics, reset where applicable, and monitoring, where included in the selected architecture."
            )
        elif starter_kind == "dol":
            doc.add_paragraph(
                f"{comm_subject} shall provide {comm_label}-capable interfaces for pump start/stop commands, run feedback, ready status, fault status, reset where applicable, and diagnostics, where included in the selected architecture."
            )
        else:
            doc.add_paragraph(
                f"{comm_subject} shall support {comm_label} communications for command, status, diagnostics, and monitoring where supported by the selected device."
            )

    doc.add_heading("6.3 Control Power and Auxiliaries", level=2)

    control_power_loads = [
        "control relays",
        "starter control circuits",
        "field instrumentation loops where required",
    ]

    if has_plc:
        control_power_loads.extend(["PLC"])

    if has_network:
        control_power_loads.append("communication interface devices where included")

    doc.add_paragraph(
        "The panel shall include control power supplies sized for "
        + ", ".join(control_power_loads[:-1])
        + " and "
        + control_power_loads[-1]
        + "."
    )

    doc.add_paragraph(
        "An emergency stop function shall remove run permission from all motor starters and place the system into a safe state."
    )
    doc.add_paragraph(
        "Door interlock and maintenance mode provisions shall be implemented as required by local regulations and end user standards."
    )

    if has_plc:
        doc.add_paragraph(
            "The panel shall include a PLC with sufficient capacity and I/O to implement the required booster control functions."
        )
    else:
        if comm_mode == "hardwired":
            doc.add_paragraph(
                "Where no PLC is included in the panel, the panel shall provide the required hardwired interfaces for connection to an external pump controller, BMS, or site control system."
            )
        else:
            doc.add_paragraph(
                "Where no PLC is included in the panel, the panel shall provide the required hardwired and/or communication interfaces for connection to an external pump controller, BMS, or site control system."
            )

    if has_scada:
        if has_plc:
            doc.add_paragraph(
                "The panel shall include the required interface for supervisory SCADA monitoring and control, including status, alarms, and commands where required."
            )
        else:
            doc.add_paragraph(
                "SCADA integration shall be provided through the external controller, BMS, or site control system connected to the panel interfaces."
            )
    else:
        if has_plc:
            doc.add_paragraph(
                "Operator control and monitoring shall be provided through local hardwired devices and PLC I/O as defined by the selected control architecture."
            )
        else:
            doc.add_paragraph(
                "Operator control and monitoring shall be provided through local hardwired devices and/or by the external controller as defined by the selected architecture."
            )

    doc.add_paragraph(
        "Pressure transmitter quantity, redundancy, scaling, and failover behavior shall be implemented according to the selected I/O configuration and project requirements."
    )

    # 7. Control Philosophy
    if starter_kind == "vfd":
        doc.add_heading("7. Control Philosophy - VSD Booster Set", level=1)
    elif starter_kind in ["ats01", "ats130"]:
        doc.add_heading("7. Control Philosophy - Soft Starter Booster Set", level=1)
    elif starter_kind == "dol_adv":
        doc.add_heading("7. Control Philosophy - DOL Advanced Booster Set", level=1)
    elif starter_kind == "dol":
        doc.add_heading("7. Control Philosophy - DOL Booster Set", level=1)
    else:
        doc.add_heading("7. Control Philosophy - Booster Set", level=1)

    philosophy_clauses = get_safe_control_philosophy_clauses(ctx, templates)
    for clause in philosophy_clauses:
        doc.add_paragraph(clause)

    # 7. Communications
    if has_plc:
        loss_comm_clause = (
            "Loss of communications shall generate an alarm and the PLC shall transition "
            "the booster set to a defined safe state according to the control philosophy."
        )
    else:
        loss_comm_clause = (
            "Loss of communications shall generate an alarm where communication monitoring "
            "is implemented, and the required fail-safe response shall be defined in the "
            "external controller or site control system according to the control philosophy."
        )

    if comm_mode == "hardwired":
        doc.add_heading("8. Communications - Hardwired", level=1)
        controller_text = (
            "PLC" if has_plc else "external controller or site control system"
        )

        doc.add_paragraph(
            f"Control and monitoring signals between the {controller_text} and the booster panel shall be hardwired via discrete and analog I/O."
        )
        doc.add_paragraph(
            "The panel shall include sufficient terminal blocks for pump commands, run feedback, fault feedback, reset signals, permissives, pressure signals, and alarm interfaces as required."
        )
        doc.add_paragraph(
            "Loss of critical hardwired permissives, such as emergency stop or pressure protection signals, shall transition the booster set to a defined safe state."
        )

    elif comm_mode == "modbus":
        doc.add_heading(f"8. Communications - {comm_label}", level=1)

        if has_plc:
            doc.add_paragraph(
                f"The PLC shall exchange command, status, and diagnostic data with the selected communication-capable devices over {comm_label}."
            )
        else:
            doc.add_paragraph(
                f"The panel shall provide {comm_label}-capable interface points or devices for connection to the external controller, BMS, or site control system, where included in the selected architecture."
            )

        doc.add_paragraph(
            "The Modbus address map, communication parameters, timeout behavior, and available data points shall be documented."
        )

    elif comm_mode == "profinet":
        doc.add_heading(f"8. Communications - {comm_label}", level=1)

        if has_plc:
            doc.add_paragraph(
                f"The PLC shall exchange cyclic command, status, and diagnostic data with the selected communication-capable devices over {comm_label}."
            )
        else:
            doc.add_paragraph(
                f"The panel shall provide {comm_label}-capable interface points or devices for connection to the external controller, BMS, or site control system, where included in the selected architecture."
            )

        doc.add_paragraph(
            "ProfiNet device names, IP addresses where applicable, GSDML files, diagnostic data, and network topology shall be documented where applicable."
        )

    else:
        doc.add_heading(f"8. Communications - {comm_label}", level=1)
        doc.add_paragraph(
            f"The selected communication interface shall support command, status, diagnostics, and monitoring over {comm_label} where included in the selected architecture."
        )
        doc.add_paragraph(
            "Addressing, communication parameters, available data points, and timeout behavior shall be documented where applicable."
        )

    if comm_mode != "hardwired":
        doc.add_paragraph(loss_comm_clause)

    # 9. FAT/SAT
    doc.add_heading("9. FAT/SAT - Minimum Acceptance Tests", level=1)

    fat_items = [
        "wiring",
        "labeling",
        "protection device settings",
        (
            "DOL starter operation"
            if starter_kind in ["dol", "dol_adv"]
            else f"{device_type} operation"
        ),
        "local controls",
        "fault indication",
    ]

    if has_plc:
        fat_items.append("automatic booster control sequences")

    if starter_kind == "dol_adv":
        fat_items.append(
            "advanced motor-management protection and diagnostic functions"
        )

    if comm_mode != "hardwired":
        fat_items.append(f"{comm_label} communication/interface tests")

    if supports_speed_reference:
        fat_items.append("speed reference and pressure control response")

    doc.add_paragraph("FAT shall verify " + ", ".join(fat_items) + ".")

    sat_items = [
        "field wiring",
        "motor rotation",
        (
            "DOL starter operation"
            if starter_kind in ["dol", "dol_adv"]
            else f"{device_type} operation"
        ),
        "protection trips",
        "local/remote control interface",
    ]

    if comm_mode != "hardwired":
        if has_plc:
            sat_items.append(f"{comm_label} communication with panel control devices")
        else:
            sat_items.append(
                f"{comm_label} interface availability and communication with the external control system where provided"
            )

    if supports_speed_reference:
        sat_items.extend(["pressure transmitter scaling", "pressure control stability"])

    if starter_kind == "dol_adv":
        sat_items.append("advanced motor-management alarm and trip verification")

    doc.add_paragraph("SAT shall verify " + ", ".join(sat_items) + ".")

    output = io.BytesIO()
    doc.save(output)

    output.seek(0)
    return output.read()
