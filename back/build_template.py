from docx import Document

doc = Document()

# Add a title
doc.add_heading('Engineering Specification Document', 0)

# Add standard intro paragraph
doc.add_paragraph('Project Configuration DNA: {{ config_id }}')
doc.add_heading('1. Overview', level=1)
doc.add_paragraph('This document outlines the standard engineering parameters for a {{ load_count }}-pump Water Booster set, rated at {{ motor_power_kw }} kW per motor.')

# Add Control Section
doc.add_heading('2. Control Panel & Enclosure', level=1)
doc.add_paragraph('The control panel shall be housed in an enclosure dimensioned at {{ enclosure_dims }} with a mounting type of {{ enclosure_mounting }}. Default cabinet rating is IP54.')

# Add BOM table
doc.add_heading('3. Key Components BOM', level=1)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Item'
hdr_cells[1].text = 'Part Number'
hdr_cells[2].text = 'Quantity'

row_cells = table.add_row().cells
row_cells[0].text = '{% tr for comp in components %}{{ comp.item_category }}'
row_cells[1].text = '{{ comp.part_number }}'
row_cells[2].text = '{{ comp.qty }}{% tr endfor %}'

# Save the template to the required folder
doc.save('app/generators/templates/Spec_Template.docx')
print('Template successfully generated.')
